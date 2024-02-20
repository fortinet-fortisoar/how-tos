import logging
import os
from constants import *
from template import *
from docx import Document
from docx.shared import Pt


class Specs:
    def __init__(self, sp_dir_path, sp_docs_path, info_json_data, outbreak_alert_data, threat_hunt_rules_data):
        self.sp_dir_path = sp_dir_path
        self.sp_docs_path = sp_docs_path
        self.info_json_data = info_json_data
        self.outbreak_alert_data = outbreak_alert_data
        self.threat_hunt_rules_data = threat_hunt_rules_data
        
    def create_specs_file_data(self):
        SPECS_FILE_NAME = f"Outbreak Response - {self.outbreak_alert_data['title']}.docx"
        specs_file = self.__create_specs_file()
        logging.debug(
            "Successfully created \"{0}\" file".format(SPECS_FILE_NAME))

    def __create_specs_file(self):
        doc = Document()
    
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        title = self.outbreak_alert_data['title']
        # Add title
        doc.add_heading(f"Outbreak Response - {title} (SPECS)", level=1)

        # Add scope section
        doc.add_heading('Scope:', level=2)
        doc.add_paragraph(scope.substitute(title=self.outbreak_alert_data['title'], link=self.outbreak_alert_data['link']))
        
        # Add scope_list section
        for line in SCOPE_LIST.splitlines():
            if line.startswith("- "):  # Handle bulleted sub-items
                paragraph = doc.add_paragraph(line[2:],style='List Bullet 2')
                # paragraph.style.paragraph_format.left_indent = docx.shared.Inches(0.5)
                # paragraph.text = 
            else:
                doc.add_paragraph(line, style='List Number')
        doc.add_page_break()

        # Add future scope section
        doc.add_heading('Future Scope:\n', level=2)
        # Add future_scope section
        for line in FUTURE_SCOPE.splitlines():
            if line.startswith("- "):  # Handle bulleted sub-items
                paragraph = doc.add_paragraph(line[2:],style='List Bullet 3')
                # paragraph.style.paragraph_format.left_indent = docx.shared.Inches(0.5)
                # paragraph.text = 
            else:
                para = doc.add_paragraph(line, style='List Number 2')
                
        doc.add_paragraph("")        
        
        # Add table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'  # Apply table grid style
        

        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Version'
        hdr_cells[3].text = 'Purpose'


        for name, type_, version, purpose in SP_CONTENTS_DATA:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = type_
            row_cells[2].text = version
            row_cells[3].text = purpose

        doc.add_page_break()
        # Add solution pack contents section
        doc.add_heading('Solution Pack Contents:', level=2)
        doc.add_paragraph("This solution pack contains the following resources:-")
        doc.add_paragraph("Outbreak Alerts Recordset", style="List Bullet") 
        # Add Record sets table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'  # Apply table grid style
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Description'
        
        # Add cells onto Record sets table
        row_cells = table.add_row().cells
        row_cells[0].text = title
        row_cells[1].text = self.outbreak_alert_data['description']
        
        # Add Threat Hunt Rules Recordset 
        doc.add_paragraph()
        doc.add_paragraph("Threat Hunt Rules Recordset", style="List Bullet") 
        
        # Add Record sets table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'  # Apply table grid style
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Rule Type'

        # Add cells onto Record sets table
        for item in self.threat_hunt_rules_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item['title']
            row_cells[1].text = self._which_rule(item['ruleType'])
        
        # Save the file
        doc.save(f"{self.sp_dir_path}/{title} (Specs).docx")
        
    def _which_rule(self, picklist_value: str)-> str:
        if picklist_value == SIGMA_PICKLIST_VALUE:
            return "Sigma"
        
        elif picklist_value == YARA_PICKLIST_VALUE:
            return "Yara"
        
        elif picklist_value == FORTINET_FABRIC_PICKLIST_VALUE:
            return "Fortinet Fabric"

        else:
            raise Exception(f"Unknown Picklist Value of value {picklist_value}")
                    

        
            
        